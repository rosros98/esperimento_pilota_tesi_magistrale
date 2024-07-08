import os
import docx
import csv

#Idea: Indicare man mano il range di pagine sulla quale lavorare ed estrarre i dati.
def conta_celle_chiave(tabella, parole_chiave):
    conteggio = 0
    for riga in tabella.rows:
        for cella in riga.cells:
            testo_cellula = cella.text.strip()
            for chiave in parole_chiave:
                if chiave.lower() in testo_cellula.lower():
                    conteggio += 1
    return conteggio

def estrai_celle_chiave(tabella, parole_chiave):
    contenuti_celle_chiave = []
    for riga in tabella.rows:
        for cella in riga.cells:
            testo_cellula = cella.text.strip()
            for chiave in parole_chiave:
                if testo_cellula.lower().startswith(chiave.lower()):
                    contenuti_celle_chiave.append(testo_cellula)
    return contenuti_celle_chiave

def main():
    nome_file_docx = 'richiesta-cautelare-clan-CAVA.docx'
    parole_chiave = ["Colloquio", "Intercettazione ambientale", "Conversazione n",
                     "Conversazione ambientale", "Ore", "Leggenda", "Tel n",
                     "Tel.", "Decreto n", "omissis"]
    pagina_iniziale = 82  # Cambia qui il numero di pagina iniziale
    pagina_finale = 1001  # Cambia qui il numero di pagina finale

    output_directory = 'db'
    if not os.path.exists(output_directory):
        os.makedirs(output_directory)

    doc = docx.Document(nome_file_docx)

    numero_celle_chiave = 0
    conversazioni = []
    for pagina in range(pagina_iniziale - 1, min(pagina_finale, len(doc.tables))):
        numero_celle_chiave += conta_celle_chiave(doc.tables[pagina], parole_chiave)
        conversazioni += estrai_celle_chiave(doc.tables[pagina], parole_chiave)

    print(f"Numero di celle corrispondenti alle parole chiave: {numero_celle_chiave}")

    sentimento = ["" for _ in conversazioni]
    emozione = ["" for _ in conversazioni]

    output_file_path = os.path.join(output_directory, 'database.csv')
    with open(output_file_path, mode='w', newline='', encoding='utf-8') as file_csv:
        writer = csv.writer(file_csv)
        writer.writerow(['Conversazione', 'Sentimento', 'Emozione'])
        for conversazione, sent, emo in zip(conversazioni, sentimento, emozione):
            writer.writerow([conversazione, sent, emo])

if __name__ == "__main__":
    main()
